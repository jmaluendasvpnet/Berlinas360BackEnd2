from .models import Colaboradores, Permisos, Login, Token, TipoDocumento, Roles, EnteAtencion, Empresas, Department, Modulos, Siniestro, SiniestroMedia, SiniestroLog
from django.views.decorators.http import require_http_methods
from django.contrib.auth.hashers import make_password
from django.core.files.storage import default_storage
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from django.utils.crypto import get_random_string
from django.contrib.auth import authenticate
from rest_framework.response import Response
from rest_framework.decorators import action
from rest_framework.viewsets import ViewSet
from rest_framework.views import APIView
from django.http import JsonResponse
from .envioCorreos import send_email
from django.utils import timezone
from django.conf import settings
from operator import itemgetter
from datetime import timedelta
import secrets
import base64
import string
import json
import jwt
import os
import re
import os





# views.py
import json
import os
import base64
from io import BytesIO
from django.http import FileResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from .models import Siniestro

def set_cell_shading(cell, shading_color="D3D3D3"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), shading_color)
    tcPr.append(shd)

def set_font_size_for_cell(cell, font_size=9):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)

def set_font_color_for_cell(cell, hex_color="808080"):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(hex_color)

def join_items_comas_y(items):
    length = len(items)
    if length == 0:
        return "N/A"
    elif length == 1:
        return items[0]
    elif length == 2:
        return f"{items[0]} y {items[1]}"
    else:
        return f"{', '.join(items[:-1])} y {items[-1]}"

@csrf_exempt
def generar_documento_con_imagen(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        siniestro_id = data.get('id', '')
        if not siniestro_id:
            return JsonResponse({"error": "No se proporcionó siniestro_id."}, status=400)
        try:
            siniestro_id = int(siniestro_id)
        except ValueError:
            return JsonResponse({"error": f"El siniestro_id '{siniestro_id}' no es un número válido."}, status=400)

        tipo_lesion = data.get('tipoLesion', [])
        parte_cuerpo = data.get('parteCuerpo', [])
        agente_accidente = data.get('agenteAccidente', [])
        mecanismo_accidente = data.get('mecanismoAccidente', [])
        actos_inseguros = data.get('actosInseguros', [])
        condiciones_peligrosas = data.get('condicionesPeligrosas', [])
        metodo = data.get('metodo', [])
        maquina = data.get('maquina', [])
        material = data.get('material', [])
        mano_obra = data.get('manoObra', [])
        medio_ambiente = data.get('medioAmbiente', [])
        metodo_analisis = data.get('metodoAnalisis', '')
        imagen_analisis_base64 = data.get('imagenAnalisis', '')

        factores_personales = data.get('factoresPersonales', [])
        factores_trabajo = data.get('factoresTrabajo', [])

        siniestro = get_object_or_404(Siniestro, pk=siniestro_id)

        plantilla_path = os.path.join(
            settings.MEDIA_ROOT,
            'plantillas',
            'FORMATO_INVESTIGACION_ACCIDENTE_LABORAL.docx'
        )

        original_doc = Document(plantilla_path)
        print("=== ESTRUCTURA Y CONTENIDO DEL DOCUMENTO ORIGINAL ===")
        for i, paragraph in enumerate(original_doc.paragraphs):
            print(f"Parágrafo {i}: {paragraph.text}")
        for i, table in enumerate(original_doc.tables):
            print(f"Tabla {i}:")
            for j, row in enumerate(table.rows):
                celdas = [cell.text for cell in row.cells]
                print(f"  Fila {j}: {celdas}")
        print("=======================================================")

        doc = DocxTemplate(plantilla_path)

        incidente = ''
        acdt_leve = ''
        acdt_grave = ''
        acdt_mortal = ''

        if siniestro.tipo_evento == "Incidente" and siniestro.gravedad == "Choque simple":
            incidente = 'X'
        elif siniestro.tipo_evento == "Incidente" and siniestro.gravedad in ["Heridos", "Heridos y muertos", "Muertos"]:
            acdt_mortal = 'X'
        elif siniestro.tipo_evento == "Siniestro" and siniestro.gravedad == "Choque simple":
            acdt_grave = 'X'
        elif siniestro.tipo_evento == "Siniestro" and siniestro.gravedad in ["Heridos", "Heridos y muertos", "Muertos"]:
            acdt_mortal = 'X'

        context = {
            "descripcion": siniestro.descripcion,
            "tipo_evento": siniestro.tipo_evento,
            "gravedad": siniestro.gravedad,
            "numero_victimas": siniestro.numero_victimas,
            "zona": siniestro.zona,
            "fecha_creacion": siniestro.fecha_creacion,
            "colaborador": siniestro.colaborador.nombres if siniestro.colaborador else "",
            "placa_vehiculo": siniestro.vehiculo.placa if siniestro.vehiculo else "",
            "empresa": siniestro.empresa.nombre_empresa if siniestro.empresa else "",
            "direccion_text": siniestro.direccion_text,
            "incidente": incidente,
            "acdt_leve": acdt_leve,
            "acdt_grave": acdt_grave,
            "acdt_mortal": acdt_mortal
        }

        doc.render(context)

        temp_output = BytesIO()
        doc.save(temp_output)
        temp_output.seek(0)
        tmp_docx = Document(temp_output)

        todos_seleccionados = {
            item.lower() for item in (
                tipo_lesion + parte_cuerpo + agente_accidente + mecanismo_accidente
            )
        }
        for table in tmp_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip().lower() in todos_seleccionados:
                        set_cell_shading(cell)

        causas_inmediatas_table = None
        for tbl in tmp_docx.tables:
            first_row_text = " ".join(cell.text for cell in tbl.rows[0].cells).lower()
            if ("causas inmediatas" in first_row_text
                or ("actos inseguros" in first_row_text and "condiciones peligrosas" in first_row_text)):
                causas_inmediatas_table = tbl
                break

        if causas_inmediatas_table and len(causas_inmediatas_table.rows[0].cells) >= 2:
            max_filas_causas = max(len(actos_inseguros), len(condiciones_peligrosas))
            for i in range(max_filas_causas):
                new_row = causas_inmediatas_table.add_row()
                val_acto = actos_inseguros[i] if i < len(actos_inseguros) else ""
                val_cond = condiciones_peligrosas[i] if i < len(condiciones_peligrosas) else ""
                new_row.cells[0].text = val_acto
                new_row.cells[1].text = val_cond
                set_font_size_for_cell(new_row.cells[0], font_size=6)
                set_font_size_for_cell(new_row.cells[1], font_size=6)

        gap_table = None
        for tbl in tmp_docx.tables:
            first_row_text = " ".join(cell.text for cell in tbl.rows[0].cells).lower()
            if "aspecto" in first_row_text and "relacionar las diferencias" in first_row_text:
                gap_table = tbl
                break

        def insertar_aspecto_en_una_fila(aspecto, lista_items):
            row = gap_table.add_row()
            row.cells[0].text = aspecto
            joined = join_items_comas_y(lista_items)
            row.cells[1].text = joined
            set_font_size_for_cell(row.cells[0], 6)
            set_font_color_for_cell(row.cells[0], "808080")
            set_font_size_for_cell(row.cells[1], 6)
            set_font_color_for_cell(row.cells[1], "808080")

        if gap_table and len(gap_table.rows[0].cells) >= 2:
            insertar_aspecto_en_una_fila("Método", metodo)
            insertar_aspecto_en_una_fila("Máquina", maquina)
            insertar_aspecto_en_una_fila("Material", material)
            insertar_aspecto_en_una_fila("Mano de Obra", mano_obra)
            insertar_aspecto_en_una_fila("Medio Ambiente", medio_ambiente)

        causas_basicas_table = None
        for tbl in tmp_docx.tables:
            first_row_text = " ".join(cell.text for cell in tbl.rows[0].cells).lower()
            if ("6. causas básicas del evento" in first_row_text
                or ("factores personales" in first_row_text and "factores de trabajo" in first_row_text)):
                causas_basicas_table = tbl
                break

        if causas_basicas_table and len(causas_basicas_table.rows[0].cells) >= 2:
            max_filas_causas_basicas = max(len(factores_personales), len(factores_trabajo))
            for i in range(max_filas_causas_basicas):
                new_row = causas_basicas_table.add_row()
                val_factores_personales = factores_personales[i] if i < len(factores_personales) else ""
                val_factores_trabajo = factores_trabajo[i] if i < len(factores_trabajo) else ""
                new_row.cells[0].text = val_factores_personales
                new_row.cells[1].text = val_factores_trabajo
                set_font_size_for_cell(new_row.cells[0], font_size=6)
                set_font_size_for_cell(new_row.cells[1], font_size=6)

        if imagen_analisis_base64:
            try:
                formato, imgstr = imagen_analisis_base64.split(';base64,')
                image_data = base64.b64decode(imgstr)
                image_stream = BytesIO(image_data)
                for paragraph in tmp_docx.paragraphs:
                    if "las soluciones propuestas de cada análisis" in paragraph.text.lower():
                        run = paragraph.add_run()
                        run.add_picture(image_stream, width=Mm(140))
                        break
            except:
                pass

        final_output = BytesIO()
        tmp_docx.save(final_output)
        final_output.seek(0)
        return FileResponse(final_output, as_attachment=True, filename='DocumentoConImagenes.docx')

    return JsonResponse({"error": "Método no permitido"}, status=405)

import json
from openai import OpenAI
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
cliente_openai = OpenAI(api_key=settings.OP_API_KEY)

@csrf_exempt
def analisis_recomendaciones_openai(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        prompt_usuario = data.get('prompt', '')

        if not prompt_usuario:
            return JsonResponse({"error": "No se proporcionó prompt."}, status=400)

        try:
            respuesta = cliente_openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Eres un asistente que genera análisis y recomendaciones en formato JSON, sin texto adicional, "
                            "basadas en eventos de seguridad. Devuelve la respuesta con la siguiente estructura:\n\n"
                            "{\n"
                            "  \"descripcion\": {\n"
                            "    \"titulo\": \"Descripción de los Hechos\",\n"
                            "    \"descripcion\": \"...\",\n"
                            "    \"porqueSucedio\": \"...\",\n"
                            "    \"tengamosEnCuenta\": \"...\"\n"
                            "  },\n"
                            "  \"recomendaciones\": [\n"
                            "    {\n"
                            "      \"titulo\": \"...\",\n"
                            "      \"contenido\": \"...\"\n"
                            "    },\n"
                            "    ...\n"
                            "  ]\n"
                            "}\n\n"
                            "El objeto \"recomendaciones\" puede contener la cantidad de elementos que consideres adecuada, "
                            "cada uno con \"titulo\" y \"contenido\"."
                        ),
                    },
                    {
                        "role": "user",
                        "content": prompt_usuario,
                    },
                ],
                temperature=0.7,
            )
            contenido = respuesta.choices[0].message.content.strip()

            try:
                json_result = json.loads(contenido)
                print(json_result)
                return JsonResponse(json_result, safe=False)
            except:
                return JsonResponse({"error": "La respuesta no es un JSON válido", "respuesta": contenido}, status=200)

        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    else:
        return JsonResponse({"error": "Método no permitido"}, status=405)














import os
import subprocess
import tempfile
import smtplib
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.utils import timezone
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication

from .models import Contrato, Propietario

EMAIL_ORIGEN = "jmaluendasbautista@gmail.com"
PASS_EMAIL = "akpa ecrt crgj uert"
LOGO_URL = "https://saas-cms-admin-sandbox.s3.us-west-2.amazonaws.com/sites/647e59513d04a300028afa72/assets/647e59b33d04a300028afa77/Logo_berlinas_blanco_fondo-transparente_DIGITAL.png"

@csrf_exempt
def docx_to_pdf_and_save_in_model(request):
    if request.method == 'POST' and 'file' in request.FILES and 'placa' in request.POST:
        placa = request.POST['placa']
        docx_file = request.FILES['file']

        try:
            contrato = Contrato.objects.get(placa=placa)
        except Contrato.DoesNotExist:
            return HttpResponse("No se encontró contrato para esa placa", status=404)

        try:
            temp_dir = tempfile.gettempdir()
            temp_docx_path = os.path.join(temp_dir, f'temp_docx_{placa}.docx')
            with open(temp_docx_path, 'wb') as temp_file:
                for chunk in docx_file.chunks():
                    temp_file.write(chunk)

            pdf_output_path = os.path.splitext(temp_docx_path)[0] + '.pdf'

            result = subprocess.run([
                'soffice', '--headless', '--convert-to', 'pdf',
                temp_docx_path, '--outdir', temp_dir
            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

            if result.returncode != 0:
                error_message = f"Error en LibreOffice: {result.stderr}"
                return HttpResponse(error_message, status=500)

            if not os.path.exists(pdf_output_path):
                return HttpResponse("El archivo PDF no fue creado correctamente", status=500)

            pdf_filename = f"Contrato_{placa}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            pdf_output_dir = os.path.join(settings.MEDIA_ROOT, 'contratos')
            if not os.path.exists(pdf_output_dir):
                os.makedirs(pdf_output_dir, exist_ok=True)
            final_pdf_path = os.path.join(pdf_output_dir, pdf_filename)
            os.rename(pdf_output_path, final_pdf_path)

            contrato.pdf_contrato.name = os.path.join('contratos', pdf_filename)
            contrato.firmado = True
            contrato.save()

            _enviar_pdf_a_propietarios(placa, final_pdf_path, pdf_filename)

            with open(final_pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename={pdf_filename}'
            return response

        except Exception as e:
            return HttpResponse(f"Error interno: {str(e)}", status=500)

    return HttpResponse("Solicitud inválida o parámetros faltantes", status=400)


def _enviar_pdf_a_propietarios(placa: str, pdf_path: str, pdf_filename: str):

    owners = Propietario.objects.filter(vehiculos_relations__vehiculo__placa=placa)
    if not owners.exists():
        return

    asunto = "Contrato de Vinculación Finalizado"
    
    for owner in owners:
        if not owner.correo:
            continue

        try:
            mensaje = MIMEMultipart()
            mensaje["From"] = EMAIL_ORIGEN
            mensaje["To"] = owner.correo
            mensaje["Subject"] = asunto

            cuerpo_html = f"""
            <!DOCTYPE html>
            <html lang="es">
            <head>
                <meta charset="UTF-8">
                <title>Contrato de Vinculación</title>
            </head>
            <body style="font-family: Arial, sans-serif; background-color: #f4f4f4; padding: 20px;">
                <div style="max-width: 700px; margin: 0 auto; padding: 20px; border: 1px solid #ccc; border-radius: 10px; background-color: #fff; box-shadow: 0px 2px 5px rgba(0, 0, 0, 0.1);">
                    <div style="text-align: center; background-color: #009944; border-radius: 10px;">
                        <img src="{LOGO_URL}" alt="Logo" style="max-width: 150px; display: block; margin: 20px auto;">
                    </div>
                    <div style="margin-top: 20px; line-height: 1.6; color: #000;">
                        <p>Estimado(a) <strong>{owner.nombres}</strong>,</p>
                        <p style="color: #000 !important;">
                            Nos complace informarle que el proceso de firma de su Contrato de Vinculación
                            con BERLINASTUR S.A. ha finalizado exitosamente.
                            Adjunto encontrará el documento en formato PDF.
                        </p>
                        <p style="color: #000 !important;">
                            Para mayor información, puede contactarnos respondiendo a este correo.
                        </p>
                        <hr style="border: 0; border-top: 1px solid #ccc; margin: 20px 0;">
                        <p style="color: #009944;"><strong>Cordialmente,</strong></p>
                        <p style="color: #000 !important;">
                            Departamento Jurídico<br>
                            <a style="text-decoration: none; color: #009944;" href="mailto:juridica@berlinasdelfonce.com">
                                juridica@berlinasdelfonce.com
                            </a><br>
                            Celular: <a style="text-decoration: none; color: #009944;" 
                                       href="https://api.whatsapp.com/send?phone=+573165269210">
                                       3165269210
                                    </a><br>
                            Teléfono: <a style="text-decoration: none; color: #009944;" href="">
                                          (601) 743 5050 ext. 1003
                                       </a><br>
                            Cra. 68D No.15-15 Zona Industrial Montevideo<br>
                            Bogotá D.C. - Colombia
                        </p>
                    </div>
                    <div style="margin-top: 30px; font-style: italic; font-size: 10px; color: #888; text-align: center;">
                        <p>© 2025 - Todos los derechos reservados</p>
                    </div>
                </div>
            </body>
            </html>
            """

            mensaje.attach(MIMEText(cuerpo_html, "html"))

            with open(pdf_path, "rb") as f:
                pdf_attach = MIMEApplication(f.read(), _subtype="pdf")
            pdf_attach.add_header('Content-Disposition', 'attachment', filename=pdf_filename)
            mensaje.attach(pdf_attach)

            smtp = smtplib.SMTP_SSL("smtp.gmail.com", 465)
            smtp.login(EMAIL_ORIGEN, PASS_EMAIL)
            smtp.sendmail(EMAIL_ORIGEN, owner.correo, mensaje.as_string())
            smtp.quit()

        except Exception as e:
            print(f"Error enviando correo a {owner.correo}: {str(e)}")






from django.http import JsonResponse
import os
from django.conf import settings

@csrf_exempt
def upload_excel(request):
    if request.method == 'POST':
        file_obj = request.FILES.get('file')
        if not file_obj:
            return JsonResponse({'error': 'No se recibió ningún archivo'}, status=400)
        
        ruta_plantillas = os.path.join(settings.MEDIA_ROOT, 'plantillas')
        os.makedirs(ruta_plantillas, exist_ok=True)
        destino = os.path.join(ruta_plantillas, file_obj.name)
        
        with open(destino, 'wb+') as f:
            for chunk in file_obj.chunks():
                f.write(chunk)
        
        print(f"Archivo guardado en: {destino}")
        
        return JsonResponse({
            'status': 'ok',
            'file_name': file_obj.name,
            'path': destino
        })
        
    return JsonResponse({'error': 'Método no permitido'}, status=405)

@csrf_exempt
def print_data(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            data = request.POST

        print("Datos recibidos:", data)

        return JsonResponse({'status': 'Datos recibidos'}, status=200)
    else:
        return JsonResponse({'error': 'Método no permitido'}, status=405)

# Funcionando en Berlinas365
def upload_image(request):
    if request.method == 'POST':
        files = request.FILES.getlist('file')
        if files:
            file_urls = []
            for file in files:
                file_path = os.path.join(settings.MEDIA_ROOT, 'face_id', file.name)
                file_path = default_storage.save(file_path, file)
                file_urls.append(file_path)
            return JsonResponse({'urls': file_urls})
        return JsonResponse({'error': 'No files uploaded'}, status=400)
    return JsonResponse({'error': 'Invalid method'}, status=405)

def text(request):
    JsonResponse({'message': 'Esta es una respuesta en texto, para la prueba de conversion a voz.'})



def create_siniestro(request):
    if request.method == 'POST':
        descripcion = request.POST.get('descripcion')
        tipo_evento = request.POST.get('tipo_evento')
        gravedad = request.POST.get('gravedad')
        numero_victimas = request.POST.get('numero_victimas')
        entes_atendieron_ids = request.POST.getlist('entes_atendieron')
        latitud = request.POST.get('latitud')
        longitud = request.POST.get('longitud')
        zona = request.POST.get('zona')
        print(descripcion, tipo_evento, gravedad, numero_victimas, entes_atendieron_ids, latitud, longitud, zona)

        if not (descripcion and tipo_evento and gravedad and numero_victimas and entes_atendieron_ids and latitud and longitud and zona):
            return JsonResponse({'error': 'Campos requeridos faltantes'}, status=400)

        siniestro = Siniestro.objects.create(
            descripcion=descripcion,
            tipo_evento=tipo_evento,
            gravedad=gravedad,
            numero_victimas=numero_victimas,
            latitud=latitud,
            longitud=longitud,
            zona=zona,
        )

        entes = EnteAtencion.objects.filter(id__in=entes_atendieron_ids)
        siniestro.entes_atendieron.set(entes)

        return JsonResponse({'siniestro_id': siniestro.id, 'message': 'Siniestro creado exitosamente'})

    return JsonResponse({'error': 'Invalid method'}, status=405)



@csrf_exempt
@require_http_methods(["POST"])
def upload_media(request):
    if request.method == 'POST':
        siniestro_id = request.POST.get('siniestro_id')
        files = request.FILES.getlist('file')
        print(files)

        if not siniestro_id or not files:
            return JsonResponse({'error': 'Siniestro ID y archivos son requeridos'}, status=400)

        try:
            siniestro = Siniestro.objects.get(id=siniestro_id)
        except Siniestro.DoesNotExist:
            return JsonResponse({'error': 'Siniestro no encontrado'}, status=404)

        file_urls = []
        for file in files:
            filename = default_storage.save(f'siniestros_media/{file.name}', file)
            file_url = default_storage.url(filename)
            file_type = file.content_type

            media = SiniestroMedia.objects.create(
                siniestro=siniestro,
                file_url=filename,
                tipo=file_type
            )

            SiniestroLog.objects.create(
                siniestro=siniestro,
                user=siniestro.colaborador,
                action=f'Medio subido: {file.name}',
                media=media
            )

            file_urls.append(file_url)
        return JsonResponse({'urls': file_urls})
    return JsonResponse({'error': 'Método inválido'}, status=405)





def subir_foto(request):
    imagen = request.FILES['imagen']
    destination = os.path.join('static', imagen.name)

    # Guardar la imagen en una carpeta específica
    with open(destination, 'wb+') as destination_file:
        for chunk in imagen.chunks():
            destination.write(chunk)

    # Aca debo poner la logica de subida de fotos

    return JsonResponse({'mensaje': 'Imagen subida correctamente'})


def generate_token():
    characters = string.ascii_uppercase + string.digits
    token = ''.join(secrets.choice(characters) for i in range(6))
    return token

import hashlib
def generate_uidb64(num_documento):
    user = Login.objects.get(documento_num=num_documento)
    print(user.last_name)
    user_id_str = str(user.documento_num.num_documento)
    timestamp = str(int(timezone.now().timestamp()))
    num_documento = str(user.documento_num.num_documento)
    data_to_hash = f"{user_id_str}{timestamp}{num_documento}"
    # Crear el hash unico
    hash_object = hashlib.sha256(data_to_hash.encode())
    uidb64 = base64.urlsafe_b64encode(
        hash_object.digest()).decode().replace('=', '')
    return uidb64


@method_decorator(csrf_exempt, name='dispatch')
class ResetPass(ViewSet):
    def send_reset_email(self, request):
        email = request.data["email"]
        dni = request.data["dni"]
        print(email, dni)
        user = Login.objects.get(email=email, documento_num=dni)
        if user:
            token = generate_token()
            first_name = user.first_name
            last_name = user.last_name
            uidb64 = generate_uidb64(user.documento_num)
            send_email(email, first_name, last_name, uidb64, token)
            Token.objects.create(
                token=token, documento_num=dni, documento_num_cryp=uidb64)
            return Response({'message': 'EMail enviado'})

    def new_pass(self, request):
        uidb64 = request.data["uidb64"]
        token = request.data["token"]
        password = request.data["password"]
        password2 = request.data["password2"]
        tok = Token.objects.get(
            token=token, documento_num_cryp=uidb64, vencido=False)
        if tok:
            user = Login.objects.get(documento_num=tok.documento_num)
            make_pass = make_password(password)
            user.password = make_pass
            user.save()
            tok.vencido = True
            tok.save()
            return JsonResponse({'message': 'Contraseña cambiada con éxito'})
        return JsonResponse({'error': 'Hubo un problema al cambiar la contraseña'}, status=500)


class LoginView(APIView):
    # Vista para el inicio de sesion con Face Id
    def post(self, request):
        username = request.data["username"]
        user = authenticate(
            username=username, password=request.data["password"])
        print(user)
        if user:
            hora_actual = timezone.now()
            user_login = Login.objects.get(username=username)
            user_login.last_login = hora_actual
            user_login.save()

            payload = {"username": user.username,
                       "nombre": user.first_name,
                       "apellido": user.last_name,
                       "rol_id": user.documento_num.rol_id.id_rol}
            token = jwt.encode(payload, 'your-secret-key',
                               algorithm='HS256')  # Genera un JWT
            
            return Response({"token": token,
                             "num_documento": user.documento_num.num_documento,
                             "username": user.username,
                             "nombre": user.first_name,
                             "apellido": user.last_name,
                             "rol_id": user.documento_num.rol_id.id_rol,
                             "message": user.first_name + " has iniciado sesion"
                             })
        else:
            return Response({"message": "Credenciales inválidas"}, status=400)


class LoginFaceView(APIView):
    def post(self, request):
        documento_num = int(request.data["documento_num"])
        print(type(documento_num))
        user = Login.objects.get(documento_num=documento_num)
        usernamee = user.username
        print(usernamee)
        userpass = user.password
        print(type(userpass))
        userLogin = authenticate(
            username=usernamee, password=userpass)
        print(userLogin)
        if userLogin:
            payload = {"username": userLogin.username,
                       "nombre": userLogin.first_name,
                       "apellido": userLogin.last_name,
                       "rol_id": userLogin.documento_num.rol_id.id_rol}
            token = jwt.encode(payload, 'your-secret-key',
                               algorithm='HS256')  # Genera un JWT

            return Response({"token": token,
                             "num_documento": userLogin.documento_num.num_documento,
                             "username": userLogin.username,
                             "nombre": userLogin.first_name,
                             "apellido": userLogin.last_name,
                             "rol_id": userLogin.documento_num.rol_id.id_rol,
                             "message": userLogin.first_name + " has iniciado sesion de manera facial"
                             })
        else:
            return Response({"message": "Credenciales inválidas"}, status=400)




class UsersView(ViewSet):
    @action(detail=True, methods=['POST'])
    def post(self, request, format=None):
        id_user = request.POST.get('id_user')
        username = request.POST.get('usuario')
        contrasena = request.POST.get('contrasena')
        passs = make_password(contrasena)
        superuserValue = request.POST.get('superuser')
        superuser = superuserValue == 'on'
        colaborador = Colaboradores.objects.get(num_documento=id_user)
        first_name = colaborador.nombres
        last_name = colaborador.apellidos
        email = colaborador.email

        Login.objects.create(password=passs, is_superuser=superuser, username=username, first_name=first_name,
                             last_name=last_name, email=email, is_staff=True, is_active=True, documento_num=colaborador)
        response_data = {'mensaje': 'Colaborador creado con éxito'}
        return JsonResponse(response_data)

    @action(detail=False, methods=['POST'])
    def create_temp_user(self, request):
        from urllib.parse import quote

        documento_num = request.POST.get('documento_num')
        empresa_id = request.POST.get('empresa')
        redirect = request.POST.get('redirect')
        print(documento_num, empresa_id, redirect)
        if not documento_num:
            return JsonResponse({'error': 'El número de documento es requerido.'}, status=400)
        
        empresa = Empresas.objects.get(id=empresa_id)
        print(empresa)

        colaborador, created = Colaboradores.objects.get_or_create(
            num_documento=documento_num,
            defaults={
                'nombres': '',
                'apellidos': '',
                'email': '',
                'tipo_documento': TipoDocumento.objects.first(),
                'rol': None,
                'empresa': empresa,
            }
        )

        def generate_random_string(length=8):
            return get_random_string(length)

        username = 'temp_' + generate_random_string(6)
        password = generate_random_string(10)
        hashed_password = make_password(password)

        expiration_time = timezone.now() + timedelta(hours=2)

        temp_role_name = 'temp_role_' + generate_random_string(6)
        temp_role = Roles.objects.create(
            detalle_rol=temp_role_name,
            empresa=empresa,
            department=Department.objects.first()
        )

        temp_user = Login.objects.create(
            username=username,
            password=hashed_password,
            is_active=True,
            is_staff=False,
            is_superuser=False,
            is_temporary=True,
            expiration_time=expiration_time,
            has_logged_in=False,
            documento_num=colaborador,
        )

        colaborador.rol_id = temp_role
        colaborador.save()

        normalized_path = self.normalize_path(redirect)

        try:
            modulo = Modulos.objects.get(link=normalized_path)
        except Modulos.DoesNotExist:
            return JsonResponse({'error': f'El módulo con link {normalized_path} no existe.'}, status=404)

        permiso = Permisos.objects.create(
            rol=temp_role,
            modulo=modulo,
            estado_permiso=True
        )

        token = generate_random_string(30)
        temp_user.single_use_token = token
        temp_user.save()

        encoded_redirect = quote(redirect, safe='')
        qr_url = f'{settings.FRONTEND_URL}/user/login?redirect={encoded_redirect}&token={token}'

        response_data = {
            'username': username,
            'password': password,
            'qr_url': qr_url,
        }
        return JsonResponse(response_data)

    def normalize_path(self, redirect_path):
        if redirect_path.endswith('/'):
            redirect_path = redirect_path[:-1]
        
        pattern = r'/\d+'
        normalized_path = re.sub(pattern, '/:num_documento', redirect_path)
        return normalized_path

@method_decorator(csrf_exempt, name='dispatch')
class MenuView(ViewSet):
    @action(detail=False, methods=['GET'])
    @method_decorator(require_http_methods(["GET"]))
    def get(self, request, rol):
        try:
            raiz_modulos = Permisos.objects.filter(
                modulo_id__id_modulo_padre=None, rol_id__id_rol=rol)
            print(raiz_modulos)
            data_menu = []
            for modulo in raiz_modulos:
                print(modulo)
                submodulos = Permisos.objects.filter(
                    modulo_id__id_modulo_padre=modulo.modulo_id.id_modulo, rol_id__id_rol=rol)
                data_sub_menu = []
                for submodulo in submodulos:
                    sub_mod = submodulo.modulo_id
                    data_sub_menu.append({
                        'id_modulo': sub_mod.id_modulo,
                        'id_permiso': submodulo.id_permiso,
                        'nom_modulo': sub_mod.nom_modulo,
                        'link': sub_mod.link,
                        'url_img': sub_mod.url_img,
                        'insertar': submodulo.permiso_insertar,
                        'eliminar': submodulo.permiso_eliminar,
                        'actualizar': submodulo.permiso_actualizar,
                        'consultar': submodulo.permiso_consultar,
                        'reportes': submodulo.permiso_reportes,
                    })

                # Ordena los subitems alfabéticamente por el nombre del módulo
                data_sub_menu = sorted(
                    data_sub_menu, key=itemgetter('nom_modulo'))

                mod = modulo.modulo_id
                data_menu.append({
                    'id_modulo': mod.id_modulo,
                    'id_permiso': modulo.id_permiso,
                    'nom_modulo': mod.nom_modulo,
                    'url_img': mod.url_img,
                    'isOpen': True,

                    'subItems': data_sub_menu,
                })

            # Ordena los items alfabéticamente por el nombre del módulo raíz
            data_menu = sorted(data_menu, key=itemgetter('nom_modulo'))

            return Response(data_menu)

        except Exception as e:
            print(str(e))
            return Response({'error': 'Ocurrió un error en el servidor'}, status=500)


# backend/app/views.py
import json
import io
from datetime import datetime
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponse
import pyodbc
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font

@csrf_exempt
def RP_Consultas05(request):
    if request.method == "POST":
        try:
            if request.content_type == "application/json":
                data = json.loads(request.body)
            elif request.content_type.startswith("multipart/form-data"):
                data = request.POST.dict()
            else:
                return JsonResponse({'error': 'Formato de contenido no soportado'}, status=400)
            
            empresa = int(data.get('empresa', 0))
            fecha_inicio = data.get('startDate')
            fecha_final = data.get('endDate')
            opcion = int(data.get('Opcion', 0))
            sub_opcion = int(data.get('SubOpcion', 0))

            print(empresa, fecha_inicio, fecha_final, opcion, sub_opcion)

            rows_list = execute_stored_procedure(
                "{CALL RP_Consultas05 (?, ?, ?, ?, ?, ?, ?, ?)}",
                (empresa, fecha_inicio, fecha_final, opcion, sub_opcion, "VACIO", "VACIO", 100)
            )

            return JsonResponse({'results': rows_list})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    else:
        return JsonResponse({'error': 'Método no permitido'}, status=405)

@csrf_exempt
def XlsxRP_Consultas05(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            results = json.loads(data['results'])
            opcion = int(data.get('Opcion'))
            sub_opcion = int(data.get('SubOpcion'))
            empresa = int(data.get('empresa'))
            start_date_str = data.get('startDate')
            start_date = datetime.fromisoformat(start_date_str.rstrip('Z'))
            year = start_date.year
            month = start_date.month

            # Información de la empresa
            empresa_info = get_empresa_info(empresa)

            # Crear libro de Excel
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Reporte"

            # Agregar datos de la empresa
            sheet['A1'] = "Nombre de la Empresa:"
            sheet['B1'] = empresa_info['nombre']
            sheet['A2'] = "NIT:"
            sheet['B2'] = empresa_info['nit']
            sheet['A3'] = "Fecha:"
            sheet['B3'] = datetime.now().strftime('%Y-%m-%d')

            # Agregar encabezados
            if results:
                headers = results[0].keys()
                for col_num, header in enumerate(headers, 1):
                    cell = sheet.cell(row=5, column=col_num, value=header)
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal='center')

            # Agregar datos
            for row_num, row_data in enumerate(results, 6):
                for col_num, (key, value) in enumerate(row_data.items(), 1):
                    sheet.cell(row=row_num, column=col_num, value=value)

            # Estilizar el Excel
            for column in sheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                sheet.column_dimensions[column_letter].width = adjusted_width

            # Guardar el Excel en un buffer
            buffer = io.BytesIO()
            workbook.save(buffer)
            buffer.seek(0)

            # Crear respuesta HTTP con el archivo Excel
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename=Reporte_{empresa_info["nombre"]}_{datetime.now().strftime("%Y%m%d%H%M%S")}.xlsx'
            return response

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    else:
        return JsonResponse({'error': 'Método no permitido'}, status=405)

def get_empresa_info(empresa_id):
    empresas = {
        277: {"nombre": "Berlinas del Fonce S.A.", "nit": "900.123.456-7"},
        278: {"nombre": "Berlitur S.A.S.", "nit": "900.765.432-1"},
        300: {"nombre": "Compañia Libertador S.A.", "nit": "900.111.222-3"},
        310: {"nombre": "Cartagena International Travels S.A.S. 'CIT'", "nit": "900.333.444-5"},
        320: {"nombre": "Tourline Express S.A.S.", "nit": "900.555.666-7"},
        9001: {"nombre": "Servicio Especial", "nit": "900.777.888-9"},
        # Agrega más empresas según sea necesario
    }
    return empresas.get(empresa_id, {"nombre": "Desconocida", "nit": "N/A"})

def execute_stored_procedure(stored_procedure, params):
    conn = ConexionDB()
    cursor = conn.cursor()
    try:
        cursor.execute(stored_procedure, params)
        rows = cursor.fetchall()
        columns = [column[0] for column in cursor.description]
        result = [dict(zip(columns, row)) for row in rows]
        return result
    finally:
        cursor.close()
        conn.close()

def ConexionDB():
    server = 'd1.berlinasdelfonce.com'
    # server = '172.16.0.25'
    database = 'Dynamix'
    username = 'Developer'
    password = '123456'
    try:
        conn = pyodbc.connect(
            f'DRIVER={{ODBC Driver 18 for SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password};TrustServerCertificate=yes;Encrypt=yes')
        return conn
    except pyodbc.Error as ex:
        print("Error al conectar a la base de datos:", ex)
        raise

from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.drawing.image import Image as OpenPyXLImage
from openpyxl.utils import get_column_letter
from django.http import HttpResponse
from django.utils import timezone
from io import BytesIO
import requests
import openpyxl
import json

def generate_excel_report(request):
    if request.method == 'POST':
        try:
            body_unicode = request.body.decode('utf-8')
            body_data = json.loads(body_unicode)

            empresa = body_data.get('empresa', 'Berlinas del Fonce S.A.')
            titulo = body_data.get('titulo', 'Título del Informe')
            anio = body_data.get('anio', '')
            mes = body_data.get('mes', '')
            data = body_data.get('data', [])
            columns = body_data.get('columns', [])

            if not data or not columns:
                return HttpResponse("Datos insuficientes para generar el reporte.", status=400)

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = 'Reporte'

            color_palette = { 
                'white': 'FFFFFF',
                'dark_green': '1C5632',
                'medium_green': '2E7D4F',
                'light_green': 'A8D5BA',
                'dark_gray': '5A5A5A',
                'light_gray': 'EDEDED',
            }

            title_font = Font(size=16, bold=True, color=color_palette['white'], name='Calibri')
            header_font = Font(bold=True, color=color_palette['white'], name='Calibri')
            info_font = Font(bold=True, color=color_palette['dark_green'], name='Calibri')
            data_font = Font(color=color_palette['dark_gray'], name='Calibri')

            center_alignment = Alignment(horizontal='center', vertical='center')
            left_alignment = Alignment(horizontal='left', vertical='center')
            right_alignment = Alignment(horizontal='right', vertical='center')

            fill_company = PatternFill(start_color=color_palette['dark_green'], end_color=color_palette['dark_green'], fill_type='solid')
            fill_title = PatternFill(start_color=color_palette['medium_green'], end_color=color_palette['medium_green'], fill_type='solid')
            fill_info = PatternFill(start_color=color_palette['light_green'], end_color=color_palette['light_green'], fill_type='solid')
            fill_header = PatternFill(start_color=color_palette['medium_green'], end_color=color_palette['medium_green'], fill_type='solid')
            fill_data_odd = PatternFill(start_color=color_palette['white'], end_color=color_palette['white'], fill_type='solid')
            fill_data_even = PatternFill(start_color=color_palette['light_green'], end_color=color_palette['light_green'], fill_type='solid')

            thin_border = Border(
                left=Side(border_style='thin', color=color_palette['dark_gray']),
                right=Side(border_style='thin', color=color_palette['dark_gray']),
                top=Side(border_style='thin', color=color_palette['dark_gray']),
                bottom=Side(border_style='thin', color=color_palette['dark_gray'])
            )

            num_columns = len(columns)
            if num_columns == 0:
                return HttpResponse("No hay columnas para generar el reporte.", status=400)

            start_col = 1
            end_col = num_columns
            data_start_row = 9
            header_row = data_start_row - 1

            start_cell = get_column_letter(start_col) + '1'
            end_cell = get_column_letter(end_col) + '1'
            ws.merge_cells(f'{start_cell}:{end_cell}')
            ws[start_cell] = empresa
            ws[start_cell].font = title_font
            ws[start_cell].alignment = center_alignment
            ws[start_cell].fill = fill_company
            ws.row_dimensions[1].height = 30

            start_cell = get_column_letter(start_col) + '2'
            end_cell = get_column_letter(end_col) + '2'
            ws.merge_cells(f'{start_cell}:{end_cell}')
            ws[start_cell] = titulo
            ws[start_cell].font = title_font
            ws[start_cell].alignment = center_alignment
            ws[start_cell].fill = fill_title
            ws.row_dimensions[2].height = 25

            ws.merge_cells('A4:C4')
            ws['A4'] = 'Fecha de Generación:'
            ws['A4'].font = info_font
            ws['A4'].fill = fill_info
            ws['A4'].alignment = right_alignment

            ws.merge_cells('D4:F4')
            ws['D4'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
            ws['D4'].font = data_font
            ws['D4'].fill = fill_info
            ws['D4'].alignment = left_alignment

            ws.merge_cells('A5:C5')
            ws['A5'] = 'Año:'
            ws['A5'].font = info_font
            ws['A5'].fill = fill_info
            ws['A5'].alignment = right_alignment

            ws.merge_cells('D5:F5')
            ws['D5'] = anio
            ws['D5'].font = data_font
            ws['D5'].fill = fill_info
            ws['D5'].alignment = left_alignment

            ws.merge_cells('A6:C6')
            ws['A6'] = 'Mes:'
            ws['A6'].font = info_font
            ws['A6'].fill = fill_info
            ws['A6'].alignment = right_alignment

            ws.merge_cells('D6:F6')
            ws['D6'] = mes
            ws['D6'].font = data_font
            ws['D6'].fill = fill_info
            ws['D6'].alignment = left_alignment

            if end_col >= 2:
                logo_col_num = end_col - 1
            else:
                logo_col_num = 1

            logo_cell = f'{get_column_letter(logo_col_num)}1'

            logo_url = 'https://saas-cms-admin-sandbox.s3.us-west-2.amazonaws.com/sites/647e59513d04a300028afa72/assets/647e59b33d04a300028afa77/Logo_berlinas_blanco_fondo-transparente_DIGITAL.png'
            response_logo = requests.get(logo_url)
            if response_logo.status_code == 200:
                logo_image = OpenPyXLImage(BytesIO(response_logo.content))
                logo_image.width = 142
                logo_image.height = 40
                ws.add_image(logo_image, logo_cell)
            else:
                print("No se pudo descargar el logo.")

            ws.row_dimensions[header_row].height = 20
            for col_num, col in enumerate(columns, start=1):
                cell = ws.cell(row=header_row, column=col_num)
                cell.value = col.get('title', '')
                cell.font = header_font
                cell.alignment = center_alignment
                cell.fill = fill_header
                cell.border = thin_border

            for idx, row_data in enumerate(data):
                row_num = data_start_row + idx
                for col_num, col in enumerate(columns, start=1):
                    data_index = col.get('dataIndex', '')
                    cell = ws.cell(row=row_num, column=col_num)
                    cell_value = row_data.get(data_index, '')
                    if isinstance(cell_value, (int, float)):
                        cell.value = cell_value
                        cell.number_format = '#,##0'
                        cell.alignment = right_alignment
                    else:
                        cell.value = str(cell_value)
                        cell.alignment = left_alignment
                    cell.font = data_font
                    if (row_num - data_start_row) % 2 == 0:
                        cell.fill = fill_data_even
                    else:
                        cell.fill = fill_data_odd
                    cell.border = thin_border

            for col_idx, col in enumerate(columns, start=1):
                column_letter = get_column_letter(col_idx)
                
                header = col.get('title', '')
                max_length = len(str(header))
                
                data_index = col.get('dataIndex', '')
                for row_data in data:
                    cell_value = row_data.get(data_index, '')
                    cell_length = len(str(cell_value))
                    if cell_length > max_length:
                        max_length = cell_length
                
                adjusted_width = max_length + 2
                
                ws.column_dimensions[column_letter].width = adjusted_width

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            filename = f'Reporte_{titulo}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            response['Content-Disposition'] = f'attachment; filename={filename}'
            wb.save(response)
            return response
        except Exception as e:
            return HttpResponse(f"Error al generar el reporte: {str(e)}", status=500)
    else:
        return HttpResponse("Método no permitido.", status=405)











from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from .models import Vehiculos, Servicio
import json
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.drawing.image import Image as OpenPyXLImage
from openpyxl.utils import get_column_letter
import requests
from io import BytesIO
import traceback

from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from .models import Vehiculos, Servicio
import json
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.drawing.image import Image as OpenPyXLImage
from openpyxl.utils import get_column_letter
import requests
from io import BytesIO
import traceback

@csrf_exempt
def generate_excel_report2(request):
    if request.method != 'POST':
        return HttpResponse('Método no permitido.', status=405)
    try:
        body_data = json.loads(request.body.decode('utf-8'))
        empresa = body_data.get('empresa', 'Berlinas del Fonce S.A.')
        titulo = body_data.get('titulo', 'Título del Informe')
        anio = body_data.get('anio', '')
        mes = body_data.get('mes', '')
        data = body_data.get('data', [])
        if not data:
            return HttpResponse('Datos insuficientes para generar el reporte.', status=400)
        columns = body_data.get('columns', [])
        existing = {c.get('dataIndex') for c in columns if c.get('dataIndex')}
        for key in data[0].keys():
            if key not in existing:
                columns.append({'title': key.replace('_', ' ').upper(), 'dataIndex': key})
        if not columns:
            return HttpResponse('Datos insuficientes para generar el reporte.', status=400)
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Reporte'
        color_palette = {
            'white': 'FFFFFF',
            'dark_green': '1C5632',
            'medium_green': '2E7D4F',
            'light_green': 'A8D5BA',
            'dark_gray': '5A5A5A',
        }
        title_font = Font(size=16, bold=True, color=color_palette['white'], name='Calibri')
        header_font = Font(bold=True, color=color_palette['white'], name='Calibri')
        info_font = Font(bold=True, color=color_palette['dark_green'], name='Calibri')
        data_font = Font(color=color_palette['dark_gray'], name='Calibri')
        center_alignment = Alignment(horizontal='center', vertical='center')
        left_alignment = Alignment(horizontal='left', vertical='center')
        right_alignment = Alignment(horizontal='right', vertical='center')
        fill_company = PatternFill(start_color=color_palette['dark_green'], end_color=color_palette['dark_green'], fill_type='solid')
        fill_title = PatternFill(start_color=color_palette['medium_green'], end_color=color_palette['medium_green'], fill_type='solid')
        fill_info = PatternFill(start_color=color_palette['light_green'], end_color=color_palette['light_green'], fill_type='solid')
        fill_header = PatternFill(start_color=color_palette['medium_green'], end_color=color_palette['medium_green'], fill_type='solid')
        fill_data_odd = PatternFill(start_color=color_palette['white'], end_color=color_palette['white'], fill_type='solid')
        fill_data_even = PatternFill(start_color=color_palette['light_green'], end_color=color_palette['light_green'], fill_type='solid')
        thin_border = Border(
            left=Side(border_style='thin', color=color_palette['dark_gray']),
            right=Side(border_style='thin', color=color_palette['dark_gray']),
            top=Side(border_style='thin', color=color_palette['dark_gray']),
            bottom=Side(border_style='thin', color=color_palette['dark_gray'])
        )
        num_columns = len(columns)
        data_start_row, header_row = 9, 8
        ws.merge_cells(f'A1:{get_column_letter(num_columns)}1')
        ws['A1'] = empresa
        ws['A1'].font = title_font
        ws['A1'].alignment = center_alignment
        ws['A1'].fill = fill_company
        ws.row_dimensions[1].height = 30
        ws.merge_cells(f'A2:{get_column_letter(num_columns)}2')
        ws['A2'] = titulo
        ws['A2'].font = title_font
        ws['A2'].alignment = center_alignment
        ws['A2'].fill = fill_title
        ws.row_dimensions[2].height = 25
        ws.merge_cells('A4:C4')
        ws['A4'] = 'Fecha de Generación:'
        ws['A4'].font = info_font
        ws['A4'].fill = fill_info
        ws['A4'].alignment = right_alignment
        ws.merge_cells('D4:F4')
        ws['D4'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        ws['D4'].font = data_font
        ws['D4'].fill = fill_info
        ws['D4'].alignment = left_alignment
        ws.merge_cells('A5:C5')
        ws['A5'] = 'Año:'
        ws['A5'].font = info_font
        ws['A5'].fill = fill_info
        ws['A5'].alignment = right_alignment
        ws.merge_cells('D5:F5')
        ws['D5'] = anio
        ws['D5'].font = data_font
        ws['D5'].fill = fill_info
        ws['D5'].alignment = left_alignment
        ws.merge_cells('A6:C6')
        ws['A6'] = 'Mes:'
        ws['A6'].font = info_font
        ws['A6'].fill = fill_info
        ws['A6'].alignment = right_alignment
        ws.merge_cells('D6:F6')
        ws['D6'] = mes
        ws['D6'].font = data_font
        ws['D6'].fill = fill_info
        ws['D6'].alignment = left_alignment
        logo_cell_anchor = f'{get_column_letter(num_columns - 1 if num_columns >= 2 else 1)}1'
        logo_url = 'https://saas-cms-admin-sandbox.s3.us-west-2.amazonaws.com/sites/647e59513d04a300028afa72/assets/647e59b33d04a300028afa77/Logo_berlinas_blanco_fondo-transparente_DIGITAL.png'
        try:
            resp_logo = requests.get(logo_url, timeout=10)
            if resp_logo.status_code == 200:
                img = OpenPyXLImage(BytesIO(resp_logo.content))
                img.width, img.height = 142, 40
                ws.add_image(img, logo_cell_anchor)
        except requests.exceptions.RequestException:
            pass
        ws.row_dimensions[header_row].height = 20
        for col_num, col_cfg in enumerate(columns, 1):
            c = ws.cell(row=header_row, column=col_num)
            c.value = col_cfg.get('title', '')
            c.font = header_font
            c.alignment = center_alignment
            c.fill = fill_header
            c.border = thin_border
        for idx, row in enumerate(data):
            r = data_start_row + idx
            for col_num, col_cfg in enumerate(columns, 1):
                key = col_cfg.get('dataIndex', '')
                cell = ws.cell(row=r, column=col_num)
                val = row.get(key, '')
                if isinstance(val, (int, float)):
                    cell.value = val
                    cell.number_format = '#,##0'
                    cell.alignment = right_alignment
                else:
                    cell.value = str(val or '')
                    cell.alignment = left_alignment
                cell.font = data_font
                cell.fill = fill_data_even if (r - data_start_row) % 2 == 0 else fill_data_odd
                cell.border = thin_border
        for col_num, col_cfg in enumerate(columns, 1):
            letter = get_column_letter(col_num)
            header_len = len(str(col_cfg.get('title', '')))
            max_len = header_len
            key = col_cfg.get('dataIndex', '')
            for row in data:
                l = len(str(row.get(key, '') or ''))
                if l > max_len:
                    max_len = l
            ws.column_dimensions[letter].width = max_len + 2
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="Reporte_{titulo.replace(" ", "_")}_{timezone.now().strftime("%Y%m%d_%H%M%S")}.xlsx"'
        wb.save(response)
        return response
    except Exception as e:
        return HttpResponse(f'Error al generar el reporte: {e}', status=500)

def export_vehiculos_data(request):
    try:
        qs = Vehiculos.objects.select_related(
            'empresa', 'marca', 'tipoLinea', 'clase', 'carroceria', 'combustible', 'ciudadBase', 'color'
        ).prefetch_related(
            'servicio__empresaOficial', 'servicio__empresaAdministra', 'servicio__tipoOperacion',
            'servicio__nivelServicio', 'servicio__servicio',
            'propietarios_relations__propietario__tipoDocumento',
            'soat_docs__aseguradora', 'revisiones_tecnomecanicas',
            'tarjetas_operacion', 'licencias_transito_docs',
            'polizas_contractuales__aseguradora', 'polizas_extracontractuales__aseguradora',
            'polizas_todo_riesgo__aseguradora'
        )
        for p, v in request.GET.items():
            if p.endswith('__in') and v:
                qs = qs.filter(**{p: v.split(',')})
        if request.GET.get('placa__icontains'):
            qs = qs.filter(placa__icontains=request.GET['placa__icontains'])
        if request.GET.get('servicio__numeroInterno'):
            qs = qs.filter(servicio__numeroInterno=request.GET['servicio__numeroInterno'])
        if request.GET.get('propietarios_relations__propietario__identificacion'):
            qs = qs.filter(propietarios_relations__propietario__identificacion=request.GET['propietarios_relations__propietario__identificacion'])
        ordering = request.GET.get('ordering')
        if ordering:
            map_order = {
                'numero_interno': 'servicio__numeroInterno',
                'empresaNombre': 'empresa__nombre_empresa',
                'marcaNombre': 'marca__nombre',
                'ciudadBaseNombre': 'ciudadBase__nombre',
            }
            qs = qs.order_by(*[(('-' if f.startswith('-') else '') + map_order.get(f.lstrip('-'), f.lstrip('-'))) for f in ordering.split(',')])
        else:
            qs = qs.order_by('placa')
        data = []
        for v in qs.distinct():
            s = getattr(v, 'servicio', None)
            servicio_info = {
                'numero_interno': s.numeroInterno if s else None,
                'empresa_oficial_nombre': s.empresaOficial.nombre_empresa if s and s.empresaOficial else None,
                'empresa_administra_nombre': s.empresaAdministra.nombre_empresa if s and s.empresaAdministra else None,
                'tipo_operacion_nombre': s.tipoOperacion.nombre if s and s.tipoOperacion else None,
                'nivel_servicio_nombre': s.nivelServicio.nombre if s and s.nivelServicio else None,
                'categoria_servicio_nombre': s.servicio.nombre if s and s.servicio else None,
                'fecha_ingreso_servicio': s.fechaIngreso.strftime('%Y-%m-%d') if s and s.fechaIngreso else None,
                'fecha_fin_servicio': s.fechaFinServicio.strftime('%Y-%m-%d') if s and s.fechaFinServicio else None,
            }
            rels = list(v.propietarios_relations.all()[:2])
            propietario_fields = {}
            for idx in range(2):
                pre = f'propietario{idx+1}_'
                if idx < len(rels):
                    r = rels[idx]
                    p = r.propietario
                    propietario_fields.update({
                        f'{pre}tipo_documento': p.tipoDocumento.denominacion if p.tipoDocumento else None,
                        f'{pre}identificacion': p.identificacion,
                        f'{pre}direccion': p.direccion,
                        f'{pre}ciudad': p.ciudad,
                        f'{pre}departamento': p.departamento,
                        f'{pre}telefono': p.telefono,
                        f'{pre}nombres': p.nombres,
                        f'{pre}apellidos': p.apellidos,
                        f'{pre}fecha_ingreso': p.fechaIngreso.strftime("%Y-%m-%d") if p.fechaIngreso else None,
                        f'{pre}correo': p.correo,
                        f'{pre}porcentaje': r.porcentaje,
                    })
                else:
                    propietario_fields.update({
                        f'{pre}tipo_documento': None,
                        f'{pre}identificacion': None,
                        f'{pre}direccion': None,
                        f'{pre}ciudad': None,
                        f'{pre}departamento': None,
                        f'{pre}telefono': None,
                        f'{pre}nombres': None,
                        f'{pre}apellidos': None,
                        f'{pre}fecha_ingreso': None,
                        f'{pre}correo': None,
                        f'{pre}porcentaje': None,
                    })
            latest_soat = v.soat_docs.order_by('-vigencia_hasta').first()
            latest_rtm = v.revisiones_tecnomecanicas.order_by('-fecha_vencimiento').first()
            latest_to = v.tarjetas_operacion.order_by('-fechaFinVigencia').first()
            latest_lt = v.licencias_transito_docs.order_by('-fecha_expedicion').first()
            latest_pc = v.polizas_contractuales.order_by('-fecha_fin_vigencia').first()
            latest_pe = v.polizas_extracontractuales.order_by('-fecha_fin_vigencia').first()
            latest_ptr = v.polizas_todo_riesgo.order_by('-fecha_fin_vigencia').first()
            item = {
                'placa': v.placa,
                'modelo': v.modelo,
                'numero_ejes': v.numeroEjes,
                'cilindraje': v.cilindraje,
                'estado_vehiculo': v.estado,
                'union_temporal': 'Sí' if v.unionTemporal else 'No',
                'empresa_nombre': v.empresa.nombre_empresa if v.empresa else None,
                'marca_nombre': v.marca.nombre if v.marca else None,
                'tipolinea_nombre': v.tipoLinea.nombre if v.tipoLinea else None,
                'clasevehiculo_nombre': v.clase.nombre if v.clase else None,
                'carroceria_nombre': v.carroceria.nombre if v.carroceria else None,
                'combustible_nombre': v.combustible.nombre if v.combustible else None,
                'ciudadbase_nombre': v.ciudadBase.nombre if v.ciudadBase else None,
                'color_nombre': v.color.nombre if v.color else None,
                **servicio_info,
                **propietario_fields,
                'soat_numero_poliza': latest_soat.numero_poliza if latest_soat else None,
                'soat_aseguradora_nombre': latest_soat.aseguradora.nombre if latest_soat and latest_soat.aseguradora else None,
                'soat_vigencia_hasta': latest_soat.vigencia_hasta.strftime('%Y-%m-%d') if latest_soat and latest_soat.vigencia_hasta else None,
                'rtm_no_certificado': latest_rtm.no_certificado if latest_rtm else None,
                'rtm_entidad_expide': latest_rtm.entidad_expide_certificado if latest_rtm else None,
                'rtm_fecha_vencimiento': latest_rtm.fecha_vencimiento.strftime('%Y-%m-%d') if latest_rtm and latest_rtm.fecha_vencimiento else None,
                'to_numero': latest_to.numero if latest_to else None,
                'to_fecha_fin_vigencia': latest_to.fechaFinVigencia.strftime('%Y-%m-%d') if latest_to and latest_to.fechaFinVigencia else None,
                'lt_doc_numero': latest_lt.numero_documento if latest_lt else None,
                'lt_doc_fecha_expedicion': latest_lt.fecha_expedicion.strftime('%Y-%m-%d') if latest_lt and latest_lt.fecha_expedicion else None,
                'pc_numero_poliza': latest_pc.numero_poliza if latest_pc else None,
                'pc_aseguradora_nombre': latest_pc.aseguradora.nombre if latest_pc and latest_pc.aseguradora else None,
                'pc_fecha_fin_vigencia': latest_pc.fecha_fin_vigencia.strftime('%Y-%m-%d') if latest_pc and latest_pc.fecha_fin_vigencia else None,
                'pe_numero_poliza': latest_pe.numero_poliza if latest_pe else None,
                'pe_aseguradora_nombre': latest_pe.aseguradora.nombre if latest_pe and latest_pe.aseguradora else None,
                'pe_fecha_fin_vigencia': latest_pe.fecha_fin_vigencia.strftime('%Y-%m-%d') if latest_pe and latest_pe.fecha_fin_vigencia else None,
                'ptr_numero_poliza': latest_ptr.numero_poliza if latest_ptr else None,
                'ptr_aseguradora_nombre': latest_ptr.aseguradora.nombre if latest_ptr and latest_ptr.aseguradora else None,
                'ptr_fecha_fin_vigencia': latest_ptr.fecha_fin_vigencia.strftime('%Y-%m-%d') if latest_ptr and latest_ptr.fecha_fin_vigencia else None,
            }
            data.append(item)
        return JsonResponse({'data': data})
    except Exception as e:
        return JsonResponse({'error': str(e), 'trace': traceback.format_exc() if settings.DEBUG else 'Error processing request.'}, status=500)


def export_vehiculos_data(request):
    try:
        qs = Vehiculos.objects.select_related(
            'empresa', 'marca', 'tipoLinea', 'clase', 'carroceria', 'combustible', 'ciudadBase', 'color'
        ).prefetch_related(
            'servicio__empresaOficial', 'servicio__empresaAdministra', 'servicio__tipoOperacion',
            'servicio__nivelServicio', 'servicio__servicio',
            'propietarios_relations__propietario__tipoDocumento',
            'soat_docs__aseguradora', 'revisiones_tecnomecanicas',
            'tarjetas_operacion', 'licencias_transito_docs',
            'polizas_contractuales__aseguradora', 'polizas_extracontractuales__aseguradora',
            'polizas_todo_riesgo__aseguradora'
        )
        for p, v in request.GET.items():
            if p.endswith('__in') and v:
                qs = qs.filter(**{p: v.split(',')})
        if request.GET.get('placa__icontains'):
            qs = qs.filter(placa__icontains=request.GET['placa__icontains'])
        if request.GET.get('servicio__numeroInterno'):
            qs = qs.filter(servicio__numeroInterno=request.GET['servicio__numeroInterno'])
        if request.GET.get('propietarios_relations__propietario__identificacion'):
            qs = qs.filter(propietarios_relations__propietario__identificacion=request.GET['propietarios_relations__propietario__identificacion'])
        ordering = request.GET.get('ordering')
        if ordering:
            map_order = {
                'numero_interno': 'servicio__numeroInterno',
                'empresaNombre': 'empresa__nombre_empresa',
                'marcaNombre': 'marca__nombre',
                'ciudadBaseNombre': 'ciudadBase__nombre',
            }
            qs = qs.order_by(*[(('-' if f.startswith('-') else '') + map_order.get(f.lstrip('-'), f.lstrip('-'))) for f in ordering.split(',')])
        else:
            qs = qs.order_by('placa')
        data = []
        for v in qs.distinct():
            s = getattr(v, 'servicio', None)
            servicio_info = {
                'numero_interno': s.numeroInterno if s else None,
                'empresa_oficial_nombre': s.empresaOficial.nombre_empresa if s and s.empresaOficial else None,
                'empresa_administra_nombre': s.empresaAdministra.nombre_empresa if s and s.empresaAdministra else None,
                'tipo_operacion_nombre': s.tipoOperacion.nombre if s and s.tipoOperacion else None,
                'nivel_servicio_nombre': s.nivelServicio.nombre if s and s.nivelServicio else None,
                'categoria_servicio_nombre': s.servicio.nombre if s and s.servicio else None,
                'fecha_ingreso_servicio': s.fechaIngreso.strftime('%Y-%m-%d') if s and s.fechaIngreso else None,
                'fecha_fin_servicio': s.fechaFinServicio.strftime('%Y-%m-%d') if s and s.fechaFinServicio else None,
            }
            rels = list(v.propietarios_relations.all()[:2])
            propietario_fields = {}
            for idx in range(2):
                pre = f'propietario{idx+1}_'
                if idx < len(rels):
                    r = rels[idx]
                    p = r.propietario
                    propietario_fields.update({
                        f'{pre}tipo_documento': p.tipoDocumento.denominacion if p.tipoDocumento else None,
                        f'{pre}identificacion': p.identificacion,
                        f'{pre}direccion': p.direccion,
                        f'{pre}ciudad': p.ciudad,
                        f'{pre}departamento': p.departamento,
                        f'{pre}telefono': p.telefono,
                        f'{pre}nombres': p.nombres,
                        f'{pre}apellidos': p.apellidos,
                        f'{pre}fecha_ingreso': p.fechaIngreso.strftime("%Y-%m-%d") if p.fechaIngreso else None,
                        f'{pre}correo': p.correo,
                        f'{pre}porcentaje': r.porcentaje,
                    })
                else:
                    propietario_fields.update({
                        f'{pre}tipo_documento': None,
                        f'{pre}identificacion': None,
                        f'{pre}direccion': None,
                        f'{pre}ciudad': None,
                        f'{pre}departamento': None,
                        f'{pre}telefono': None,
                        f'{pre}nombres': None,
                        f'{pre}apellidos': None,
                        f'{pre}fecha_ingreso': None,
                        f'{pre}correo': None,
                        f'{pre}porcentaje': None,
                    })
            latest_soat = v.soat_docs.order_by('-vigencia_hasta').first()
            latest_rtm = v.revisiones_tecnomecanicas.order_by('-fecha_vencimiento').first()
            latest_to = v.tarjetas_operacion.order_by('-fechaFinVigencia').first()
            latest_lt = v.licencias_transito_docs.order_by('-fecha_expedicion').first()
            latest_pc = v.polizas_contractuales.order_by('-fecha_fin_vigencia').first()
            latest_pe = v.polizas_extracontractuales.order_by('-fecha_fin_vigencia').first()
            latest_ptr = v.polizas_todo_riesgo.order_by('-fecha_fin_vigencia').first()
            item = {
                'placa': v.placa,
                'modelo': v.modelo,
                'numero_ejes': v.numeroEjes,
                'cilindraje': v.cilindraje,
                'estado_vehiculo': v.estado,
                'union_temporal': 'Sí' if v.unionTemporal else 'No',
                'empresa_nombre': v.empresa.nombre_empresa if v.empresa else None,
                'marca_nombre': v.marca.nombre if v.marca else None,
                'tipolinea_nombre': v.tipoLinea.nombre if v.tipoLinea else None,
                'clasevehiculo_nombre': v.clase.nombre if v.clase else None,
                'carroceria_nombre': v.carroceria.nombre if v.carroceria else None,
                'combustible_nombre': v.combustible.nombre if v.combustible else None,
                'ciudadbase_nombre': v.ciudadBase.nombre if v.ciudadBase else None,
                'color_nombre': v.color.nombre if v.color else None,
                **servicio_info,
                **propietario_fields,
                'soat_numero_poliza': latest_soat.numero_poliza if latest_soat else None,
                'soat_aseguradora_nombre': latest_soat.aseguradora.nombre if latest_soat and latest_soat.aseguradora else None,
                'soat_vigencia_hasta': latest_soat.vigencia_hasta.strftime('%Y-%m-%d') if latest_soat and latest_soat.vigencia_hasta else None,
                'rtm_no_certificado': latest_rtm.no_certificado if latest_rtm else None,
                'rtm_entidad_expide': latest_rtm.entidad_expide_certificado if latest_rtm else None,
                'rtm_fecha_vencimiento': latest_rtm.fecha_vencimiento.strftime('%Y-%m-%d') if latest_rtm and latest_rtm.fecha_vencimiento else None,
                'to_numero': latest_to.numero if latest_to else None,
                'to_fecha_fin_vigencia': latest_to.fechaFinVigencia.strftime('%Y-%m-%d') if latest_to and latest_to.fechaFinVigencia else None,
                'lt_doc_numero': latest_lt.numero_documento if latest_lt else None,
                'lt_doc_fecha_expedicion': latest_lt.fecha_expedicion.strftime('%Y-%m-%d') if latest_lt and latest_lt.fecha_expedicion else None,
                'pc_numero_poliza': latest_pc.numero_poliza if latest_pc else None,
                'pc_aseguradora_nombre': latest_pc.aseguradora.nombre if latest_pc and latest_pc.aseguradora else None,
                'pc_fecha_fin_vigencia': latest_pc.fecha_fin_vigencia.strftime('%Y-%m-%d') if latest_pc and latest_pc.fecha_fin_vigencia else None,
                'pe_numero_poliza': latest_pe.numero_poliza if latest_pe else None,
                'pe_aseguradora_nombre': latest_pe.aseguradora.nombre if latest_pe and latest_pe.aseguradora else None,
                'pe_fecha_fin_vigencia': latest_pe.fecha_fin_vigencia.strftime('%Y-%m-%d') if latest_pe and latest_pe.fecha_fin_vigencia else None,
                'ptr_numero_poliza': latest_ptr.numero_poliza if latest_ptr else None,
                'ptr_aseguradora_nombre': latest_ptr.aseguradora.nombre if latest_ptr and latest_ptr.aseguradora else None,
                'ptr_fecha_fin_vigencia': latest_ptr.fecha_fin_vigencia.strftime('%Y-%m-%d') if latest_ptr and latest_ptr.fecha_fin_vigencia else None,
            }
            data.append(item)
        return JsonResponse({'data': data})
    except Exception as e:
        return JsonResponse({'error': str(e), 'trace': traceback.format_exc() if settings.DEBUG else 'Error processing request.'}, status=500)













import re
import os
import cv2
import threading
import matplotlib
import numpy as np
matplotlib.use('Agg')
from PIL import Image
Image.MAX_IMAGE_PIXELS = None
from paddleocr import PaddleOCR
from django.conf import settings
from rest_framework import status
from tempfile import TemporaryDirectory
from pdf2image import convert_from_path
from rest_framework.views import APIView
from rest_framework.response import Response
from django.core.files.storage import default_storage
from rest_framework.permissions import AllowAny
from rest_framework.parsers import MultiPartParser, FormParser

TARGET_WIDTH_LICENSE = 5100
TARGET_HEIGHT_LICENSE = 6600
TARGET_WIDTH_POLICY = 2480
TARGET_HEIGHT_POLICY = 3508

ROIS = {
    'soat': {
        'bolivar': { "fecha_expedicion": (610, 1150, 1060, 1250), "vigencia_desde": (1280, 1150, 1690, 1250), "vigencia_hasta": (1950, 1150, 2360, 1250), "numero_poliza": (450, 1740, 1020, 1862), "placa": (1090, 1740, 1458, 1862), },
        'aseguradora': { "fecha_expedicion": (410, 1010, 950, 1130), "vigencia_desde": (1110, 1010, 1620, 1130), "vigencia_hasta": (1850, 1010, 2370, 1130), "numero_poliza": (300, 1630, 910, 1810), "placa": (980, 1640, 1480, 1800), },
        'mundial': { "fecha_expedicion": (410, 1010, 950, 1130), "vigencia_desde": (1110, 1010, 1620, 1130), "vigencia_hasta": (1850, 1010, 2370, 1130), "numero_poliza": (300, 1630, 910, 1810), "placa": (980, 1640, 1480, 1800), },
        'estado': { "fecha_expedicion": (410, 1050, 950, 1180), "vigencia_desde": (1150, 1050, 1620, 1180), "vigencia_hasta": (1890, 1050, 2370, 1180), "numero_poliza": (280, 1630, 910, 1810), "placa": (1160, 1640, 1530, 1800), },
        'unknown': { "fecha_expedicion": (610, 1150, 1060, 1250), "vigencia_desde": (1280, 1150, 1690, 1250), "vigencia_hasta": (1950, 1150, 2360, 1250), "numero_poliza": (450, 1740, 1020, 1862), "placa": (1090, 1740, 1458, 1862), }
    },
    'tecno': {
        'unknown': { "entidad_expide_certificado": (1920, 2600, 4490, 2750), "nit": (1300, 2870, 2200, 2999), "fecha_expedicion": (1380, 3080, 2200, 3270), "no_certificado": (2100, 2000, 3000, 2150), "fecha_vencimiento": (3350, 3100, 4300, 3250), "placa": (1100, 3730, 2300, 3870) }
    },
    'operacion': {
        'unknown': { "no_certificado": (2200, 1570, 3000, 1750), "placa": (1530, 2100, 2300, 2250), "nit": (1530, 4270, 2300, 4430), "fecha_expedicion": (1530, 4800, 2300, 4950), "vigencia_desde": (2000, 5000, 2600, 5150), "fecha_vencimiento": (3100, 5000, 3650, 5150), }
    }
}

INSURER_KEYWORDS = {
    'bolivar': ['BOLIVAR'],
    'estado': ['ESTADO S.A.'],
    'mundial': ['MUNDIAL', 'ASEGURADORA'],
}

POLICY_TYPE_KEYWORDS = {
    'contractual': ['RESPONSABILIDAD CIVIL CONTRACTUAL'],
    'extracontractual': ['RESPONSABILIDAD CIVIL EXTRACONTRACTUAL'],
}

LICENSE_TECNO_KEYWORDS = ['CERTIFICADO DE REVISION TÉCNICO MECANICA', 'TECNO MECANICA', 'TÉCNICO MECANICA', 'DATOS CENTRO DIAGNOSTICO', 'CERTIFICADO DE REVISION TECNICO MECANICA Y DE EMISIONES CONTAMINANTES', 'TECNICO MECANICA']
LICENSE_SOAT_KEYWORDS = ['SOAT', 'PRIMA SOAT', 'ASEGuRADORY']
LICENSE_OPERACION_KEYWORDS = ['TARJETA DE OPERACIóN', 'TARJETA DE']

try:
    ocr_instance = PaddleOCR(
        use_angle_cls=True,
        lang='es',
        det_db_box_thresh=0.5,
        use_gpu=False,
        det_model_dir='./paddleocr_models/det_lite',
        rec_model_dir='./paddleocr_models/rec_lite',
        show_log=False
    )
except Exception as e:
    print(f"Error initializing PaddleOCR with GPU: {e}. Falling back to CPU.")
    ocr_instance = PaddleOCR(
        use_angle_cls=True,
        lang='es',
        det_db_box_thresh=0.5,
        use_gpu=False,
        det_model_dir='./paddleocr_models/det_lite',
        rec_model_dir='./paddleocr_models/rec_lite',
        show_log=False
    )

ocr_lock = threading.Lock()

def resize_image(image, width, height):
    return cv2.resize(image, (width, height), interpolation=cv2.INTER_AREA)

def preprocess_license_image(image):
    if image is None:
        raise ValueError("La imagen de entrada es None.")
    resized = resize_image(image, TARGET_WIDTH_LICENSE, TARGET_HEIGHT_LICENSE)
    gray = cv2.cvtColor(resized, cv2.COLOR_BGR2GRAY)
    _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
    return cv2.cvtColor(thresh, cv2.COLOR_GRAY2BGR)

def preprocess_policy_image(image):
    if image is None:
        raise ValueError("La imagen de entrada es None.")
    resized = resize_image(image, TARGET_WIDTH_POLICY, TARGET_HEIGHT_POLICY)
    gray = cv2.cvtColor(resized, cv2.COLOR_BGR2GRAY)
    _, thresh = cv2.threshold(gray, 180, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return cv2.cvtColor(thresh, cv2.COLOR_GRAY2BGR)

def detect_insurer(extracted_text):
    text_upper = extracted_text.upper()
    for insurer, keywords in INSURER_KEYWORDS.items():
        for keyword in keywords:
            if keyword.upper() in text_upper:
                return insurer
    return 'unknown'

def detect_policy_type_from_text_internal(extracted_text):
    text_upper = extracted_text.upper()
    for policy_type, keywords_list in POLICY_TYPE_KEYWORDS.items():
        for keyword in keywords_list:
            if keyword.upper() in text_upper:
                if "SERVICIO PUBLICO PASAJEROS" in text_upper:
                    if policy_type == "contractual":
                        return 'contractual'
                    if policy_type == "extracontractual":
                        return 'extracontractual'
                else:
                    return policy_type
    return 'desconocido'

def combined_detect_document_type(extracted_text):
    text_upper = extracted_text.upper()
    doc_type = 'desconocido'
    insurer = 'unknown'

    policy_specific_type = detect_policy_type_from_text_internal(extracted_text)
    if policy_specific_type != 'desconocido':
        doc_type = policy_specific_type
        return doc_type, insurer

    if any(keyword.upper() in text_upper for keyword in LICENSE_TECNO_KEYWORDS):
        doc_type = 'tecno'
    elif any(keyword.upper() in text_upper for keyword in LICENSE_SOAT_KEYWORDS) or ('SOAT' in text_upper and any(ins_kw.upper() in text_upper for ins_kw_list in INSURER_KEYWORDS.values() for ins_kw in ins_kw_list)):
        doc_type = 'soat'
        insurer = detect_insurer(extracted_text)
    elif any(keyword.upper() in text_upper for keyword in LICENSE_OPERACION_KEYWORDS):
        doc_type = 'operacion'
            
    return doc_type, insurer

def is_box_center_in_roi(box_coords, roi_coords):
    roi_x_start, roi_y_start, roi_x_end, roi_y_end = roi_coords
    box_center_x = sum(p[0] for p in box_coords) / 4
    box_center_y = sum(p[1] for p in box_coords) / 4
    return (roi_x_start <= box_center_x <= roi_x_end and
            roi_y_start <= box_center_y <= roi_y_end)

def extract_data_from_full_ocr_license(full_ocr_result, rois_map):
    extracted_data = {field: [] for field in rois_map.keys()}
    if full_ocr_result and full_ocr_result[0]:
        for line_info in full_ocr_result[0]:
            box_coords = line_info[0]
            text_detected = line_info[1][0]
            for field, roi_coords in rois_map.items():
                if is_box_center_in_roi(box_coords, roi_coords):
                    extracted_data[field].append(text_detected)
    final_data = {}
    for field, texts in extracted_data.items():
        final_data[field] = " ".join(texts).strip() if texts else None
    return final_data

def extract_policy_details_from_text_internal(text_content):
    data = {
        "placa": None,
        "vigencia_desde": None,
        "vigencia_hasta": None,
        "numero_poliza": None,
        "fecha_expedicion": None,
    }

    # --- INICIO DE LA SECCIÓN MODIFICADA ---
    def convertir_fecha_a_ddmmyyyy(fecha_texto):
        """Convierte una fecha en formato 'DD de Mes de AAAA' a 'DD/MM/AAAA'."""
        if not fecha_texto:
            return None
            
        meses = {
            'enero': '01', 'febrero': '02', 'marzo': '03', 'abril': '04',
            'mayo': '05', 'junio': '06', 'julio': '07', 'agosto': '08',
            'septiembre': '09', 'octubre': '10', 'noviembre': '11', 'diciembre': '12'
        }
        
        try:
            partes = fecha_texto.lower().split(' de ')
            if len(partes) == 3:
                dia = partes[0].strip().zfill(2)
                nombre_mes = partes[1].strip()
                anio = partes[2].strip()
                
                numero_mes = meses.get(nombre_mes)
                
                if numero_mes and anio.isdigit() and len(anio) == 4:
                    return f"{dia}/{numero_mes}/{anio}"
        except Exception:
            # Si hay algún error, devuelve el texto original para no perder el dato.
            return fecha_texto
        
        # Si no se pudo convertir, devuelve el texto original.
        return fecha_texto

    # --- FIN DE LA SECCIÓN MODIFICADA ---

    placa_match = re.search(r"PLACA:\s*([A-Z0-9]{3,7})", text_content, re.IGNORECASE)
    if placa_match:
        data["placa"] = placa_match.group(1).strip()
    else:
        placa_match_alternative = re.search(r"\b([A-Z]{3}\s?[0-9]{3})\b", text_content)
        if placa_match_alternative:
            data["placa"] = placa_match_alternative.group(1).replace(" ", "")
        else:
            match_near_tomador = re.search(r"([A-Z]{3}\s?[0-9]{3})\s*\n\s*TOMADOR:", text_content, re.IGNORECASE)
            if match_near_tomador:
                data["placa"] = match_near_tomador.group(1).replace(" ", "")

    vigencia_match = re.search(r"vigencias desde\s*(\d{1,2}/\d{1,2}/\d{4})\s*hasta\s*(\d{1,2}/\d{1,2}/\d{4})", text_content, re.IGNORECASE)
    if vigencia_match:
        data["vigencia_desde"] = vigencia_match.group(1).strip()
        data["vigencia_hasta"] = vigencia_match.group(2).strip()
    
    poliza_match = re.search(r"número de póliza\s*([A-Z0-9\s-]+)\s*para las vigencias", text_content, re.IGNORECASE)
    if poliza_match:
        data["numero_poliza"] = poliza_match.group(1).strip()
    else:
        poliza_match_alt = re.search(r"póliza\s*([A-Z0-9\s-]+)\s*para", text_content, re.IGNORECASE)
        if poliza_match_alt:
            data["numero_poliza"] = poliza_match_alt.group(1).strip()
        else:
            poliza_match_alt_2 = re.search(r"póliza\s*([A-Z]{2,3}\s?[0-9]+)", text_content, re.IGNORECASE)
            if poliza_match_alt_2:
                data["numero_poliza"] = poliza_match_alt_2.group(1).replace(" ","").strip()

    # Regex para buscar la fecha de expedición con el formato "DD de Mes de AAAA"
    expedicion_match = re.search(r"expide en.*?el\s*(\d{1,2}\s*de\s*[A-Za-z]+\s*de\s*\d{4})", text_content, re.IGNORECASE)
    if expedicion_match:
        # Normaliza espacios multiples en la fecha encontrada
        fecha_texto_original = re.sub(r'\s+', ' ', expedicion_match.group(1)).strip()
        data["fecha_expedicion"] = convertir_fecha_a_ddmmyyyy(fecha_texto_original)
    else:
        # Búsqueda alternativa por si el patrón anterior falla
        expedicion_match_alt = re.search(r"BOGOTA D\.C\.\s*(\d{1,2}\s*de\s*[A-Za-z]+\s*de\s*\d{4})", text_content, re.IGNORECASE)
        if expedicion_match_alt:
            fecha_texto_original = re.sub(r'\s+', ' ', expedicion_match_alt.group(1)).strip()
            data["fecha_expedicion"] = convertir_fecha_a_ddmmyyyy(fecha_texto_original)

    return data

class UploadLicenseView(APIView):
    permission_classes = [AllowAny]
    authentication_classes = []
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, format=None):
        if 'files' not in request.FILES:
            return Response({"error": "No se encontraron archivos en la solicitud."}, status=status.HTTP_400_BAD_REQUEST)

        files = request.FILES.getlist('files')
        password = request.data.get('password', None)
        
        response_payload = {
            "soat": None, "tecno": None, "operacion": None, 
            "contractual": None, "extracontractual": None,
            "otros_documentos": [], "errores": []
        }
        processed_main_doc_types = set()

        for file_idx, file_obj in enumerate(files):
            original_file_name = file_obj.name
            saved_file_path = None
            current_file_extracted_data = {}
            rel_path = None
            
            try:
                clean_filename_base = "".join(c if c.isalnum() else '_' for c in os.path.splitext(original_file_name)[0])
                clean_filename_ext = os.path.splitext(original_file_name)[1]
                unique_filename = f"{clean_filename_base}_{file_idx}{clean_filename_ext}"
                
                saved_file_path = default_storage.save(os.path.join('licenses', unique_filename), file_obj)
                full_file_path = default_storage.path(saved_file_path)
                
                rel_path = saved_file_path
                try:
                    if hasattr(settings, 'MEDIA_ROOT') and settings.MEDIA_ROOT and settings.MEDIA_ROOT in saved_file_path:
                        rel_path = os.path.relpath(saved_file_path, settings.MEDIA_ROOT)
                except Exception:
                    pass

                current_file_extracted_data['soporte'] = rel_path
                current_file_extracted_data['archivo_original'] = original_file_name

                ext = os.path.splitext(full_file_path)[1].lower()
                image_arrays = []

                if ext == '.pdf':
                    fmt = 'png'
                    convert_kwargs = {'dpi': 300, 'fmt': fmt, 'thread_count': 4}
                    if password: convert_kwargs['userpw'] = password
                    pages = convert_from_path(full_file_path, first_page=1, last_page=1, **convert_kwargs)
                    if not pages: raise ValueError("El PDF no contiene páginas o la conversión falló.")
                    pil_image = pages[0]
                    open_cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
                    image_arrays.append(open_cv_image)
                elif ext in ['.jpg', '.jpeg', '.png']:
                    img = cv2.imread(full_file_path)
                    if img is None: raise ValueError(f"No se pudo leer la imagen: {full_file_path}")
                    image_arrays.append(img)
                else:
                    error_detail = {
                        "archivo": original_file_name,
                        "error": "Formato no soportado."
                    }
                    if rel_path: error_detail["soporte"] = rel_path
                    response_payload["errores"].append(error_detail)
                    
                    if saved_file_path and default_storage.exists(saved_file_path):
                         default_storage.delete(saved_file_path)
                    continue

                img_array_source = image_arrays[0]
                if img_array_source is None or img_array_source.size == 0:
                    raise ValueError("La imagen cargada está vacía o es inválida.")

                img_proc_for_license = preprocess_license_image(img_array_source.copy())
                
                raw_ocr_results = []
                full_text = ""

                with ocr_lock:
                    raw_ocr_results = ocr_instance.ocr(np.ascontiguousarray(img_proc_for_license), cls=True)
                
                extracted_lines = []
                if raw_ocr_results and raw_ocr_results[0] is not None:
                    for line_info in raw_ocr_results[0]:
                        if len(line_info) >= 2 and isinstance(line_info[1], (tuple, list)):
                            extracted_lines.append(str(line_info[1][0]))
                full_text_from_license_prep = "\n".join(extracted_lines)

                doc_type, insurer = combined_detect_document_type(full_text_from_license_prep)
                current_file_extracted_data['tipo_documento_detectado'] = doc_type

                if doc_type in ['soat', 'tecno', 'operacion']:
                    current_rois_map = ROIS[doc_type].get(insurer, ROIS[doc_type]['unknown'])
                    data_from_rois = extract_data_from_full_ocr_license(raw_ocr_results, current_rois_map)
                    current_file_extracted_data.update(data_from_rois)
                    if insurer != 'unknown' and doc_type == 'soat':
                        current_file_extracted_data['aseguradora_detectada'] = insurer
                
                elif doc_type in ['contractual', 'extracontractual']:
                    img_proc_for_policy = preprocess_policy_image(img_array_source.copy())
                    policy_raw_ocr_results = []
                    with ocr_lock:
                        policy_raw_ocr_results = ocr_instance.ocr(np.ascontiguousarray(img_proc_for_policy), cls=True)
                    policy_extracted_lines = []
                    if policy_raw_ocr_results and policy_raw_ocr_results[0] is not None:
                        for line_info in policy_raw_ocr_results[0]:
                            print(line_info)
                            if len(line_info) >= 2 and isinstance(line_info[1], (tuple, list)):
                                policy_extracted_lines.append(str(line_info[1][0]))
                    full_text_from_policy_prep = "\n".join(policy_extracted_lines)

                    policy_number = None
                    for line in policy_extracted_lines:
                        m = re.search(r'\b[A-Z]{2}\d{6}\b', line)
                        if m:
                            policy_number = m.group()
                            break

                    extracted_details = extract_policy_details_from_text_internal(full_text_from_policy_prep)
                    extracted_details['numero_poliza'] = policy_number
                    current_file_extracted_data.update(extracted_details)

                else: 
                    current_file_extracted_data["mensaje"] = "Tipo de documento no reconocido o información insuficiente."

                is_primary_doc_type = doc_type in response_payload and doc_type not in ["otros_documentos", "errores"]
                
                if is_primary_doc_type and doc_type not in processed_main_doc_types:
                    response_payload[doc_type] = current_file_extracted_data
                    processed_main_doc_types.add(doc_type)
                else:
                    response_payload["otros_documentos"].append(current_file_extracted_data)

            except Exception as e:
                print(f"Error detallado procesando {original_file_name}: {type(e).__name__} - {e}")
                error_entry = {
                    "archivo": original_file_name,
                    "error": f"Error al procesar: {str(e)}"
                }
                if rel_path: 
                    error_entry["soporte"] = rel_path
                elif saved_file_path:
                    try:
                        temp_rel_path = saved_file_path
                        if hasattr(settings, 'MEDIA_ROOT') and settings.MEDIA_ROOT and settings.MEDIA_ROOT in saved_file_path:
                            temp_rel_path = os.path.relpath(saved_file_path, settings.MEDIA_ROOT)
                        error_entry["soporte"] = temp_rel_path
                    except Exception:
                         error_entry["soporte"] = "Ruta de archivo no disponible tras error."
                
                response_payload["errores"].append(error_entry)

        if not response_payload["otros_documentos"]:
            del response_payload["otros_documentos"]
        if not response_payload["errores"]:
            del response_payload["errores"]
        print(response_payload)
        return Response(response_payload, status=status.HTTP_200_OK)










from uuid import uuid4
class UploadView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, format=None):
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "No se proporcionó ningún archivo."}, status=status.HTTP_400_BAD_REQUEST)

        # Generar un nombre único para el archivo
        ext = os.path.splitext(file_obj.name)[1]
        filename = f"{uuid4().hex}{ext}"

        # Especificar la subcarpeta
        subdirectory = os.path.join(settings.MEDIA_ROOT, 'uploads_eval')
        os.makedirs(subdirectory, exist_ok=True)  # Crear la carpeta si no existe

        upload_path = os.path.join(subdirectory, filename)

        # Guardar el archivo
        with open(upload_path, 'wb+') as destination:
            for chunk in file_obj.chunks():
                destination.write(chunk)

        # Obtener la URL del archivo
        file_url = os.path.join(settings.MEDIA_URL, 'uploads_eval', filename)

        return Response({"url": file_url}, status=status.HTTP_201_CREATED)
    




import pandas as pd
import uuid
from io import BytesIO
from django.http import HttpResponse
from django.core.files.storage import default_storage
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, serializers
from .models import Vehiculos, Soat, RevisionTecnomecanica, TarjetaOperacion, LicenciaTransito, PolizaContractual, PolizaExtracontractual
from .serializers import SoatSlr, RevisionTecnomecanicaSlr, TarjetaOperacionSlr, LicenciaTransitoSlr, PolizaContractualSlr, PolizaExtracontractualSlr

class BulkUploadDocsAPIView(APIView):
    def post(self, request, *args, **kwargs):
        excel_file = request.FILES.get('excel')
        if not excel_file:
            return Response({'error': 'No se encontró el archivo excel.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(excel_file)
            df = df.where(pd.notnull(df), None)
        except Exception as e:
            return Response({'error': f'Error al leer el archivo Excel: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

        doc_serializers = {
            'soat': SoatSlr,
            'revision-tecnomecanica': RevisionTecnomecanicaSlr,
            'tarjeta-operacion': TarjetaOperacionSlr,
            'licencia-transito': LicenciaTransitoSlr,
            'poliza-contractual': PolizaContractualSlr,
            'poliza-extracontractual': PolizaExtracontractualSlr
        }
        
        failed_rows_data = []
        success_count = 0

        for index, row in df.iterrows():
            row_data_for_error = row.to_dict()
            row_data_for_error = {k: v for k, v in row_data_for_error.items() if pd.notna(v)}

            try:
                tipo = row.get('tipo_documento')
                placa = row.get('placa')
                file_name = row.get('file_name')

                if not all([tipo, placa, file_name]):
                    raise serializers.ValidationError("Faltan datos esenciales (tipo_documento, placa, file_name).")

                file_obj = request.FILES.get(file_name)
                if not file_obj:
                    raise serializers.ValidationError(f"El archivo '{file_name}' no fue encontrado en la carga.")

                if tipo not in doc_serializers:
                    raise serializers.ValidationError(f"Tipo de documento '{tipo}' no es válido.")

                try:
                    vehiculo = Vehiculos.objects.get(placa=placa)
                except Vehiculos.DoesNotExist:
                    raise serializers.ValidationError(f"Vehículo con placa '{placa}' no encontrado.")
                
                filename = f'docs_xlsx/{uuid.uuid4().hex}_{file_obj.name}'
                path = default_storage.save(filename, file_obj)
                
                data = row.to_dict()
                data['soporte'] = path
                
                serializer_class = doc_serializers[tipo]
                serializer = serializer_class(data=data)

                serializer.is_valid(raise_exception=True)
                serializer.save(vehiculo=vehiculo)
                success_count += 1

            except Exception as e:
                error_detail = str(e.detail) if hasattr(e, 'detail') else str(e)
                row_data_for_error['motivo_error'] = error_detail
                failed_rows_data.append(row_data_for_error)
                continue
        
        if failed_rows_data:
            failed_df = pd.DataFrame(failed_rows_data)
            if 'motivo_error' in failed_df.columns:
                cols = ['motivo_error'] + [col for col in failed_df.columns if col != 'motivo_error']
                failed_df = failed_df[cols]

            output_buffer = BytesIO()
            with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
                failed_df.to_excel(writer, index=False, sheet_name='Errores')
            output_buffer.seek(0)

            response = HttpResponse(
                output_buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename="reporte_errores_carga.xlsx"'
            response['X-Upload-Status'] = f'{success_count} registros creados. {len(failed_rows_data)} con errores.'
            response['Access-Control-Expose-Headers'] = 'X-Upload-Status, Content-Disposition'
            response.status_code = status.HTTP_400_BAD_REQUEST
            return response

        return Response(
            {'message': f'Carga masiva completada. Se crearon {success_count} registros exitosamente.'},
            status=status.HTTP_201_CREATED
        )
    


import pandas as pd
import uuid
from io import BytesIO
from datetime import datetime
from django.http import HttpResponse
from django.core.files.storage import default_storage
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, serializers
from .models import Vehiculos, Soat, RevisionTecnomecanica, TarjetaOperacion, PolizaContractual, PolizaExtracontractual
from .serializers import SoatSlr, RevisionTecnomecanicaSlr, TarjetaOperacionSlr, PolizaContractualSlr, PolizaExtracontractualSlr

def normalize_date(value):
    if pd.isna(value) or value is None:
        return None
    try:
        dt = pd.to_datetime(value, errors='coerce')
        if pd.notna(dt):
            return dt.strftime('%Y-%m-%d')
    except Exception:
        pass
    return None

class UploadCorrectedXlsxAPIView(APIView):
    def post(self, request, *args, **kwargs):
        excel_file = request.FILES.get('excel')
        if not excel_file:
            return Response({'error': 'No se encontró el archivo excel.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(excel_file, parse_dates=False)
            df.replace("CAMPO REQUERIDO", None, inplace=True)
            df = df.where(pd.notnull(df), None)
        except Exception as e:
            return Response({'error': f'Error al leer el archivo Excel: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)
        
        doc_serializers = {
            'soat': SoatSlr,
            'tecno': RevisionTecnomecanicaSlr,
            'operacion': TarjetaOperacionSlr,
            'contractual': PolizaContractualSlr,
            'extracontractual': PolizaExtracontractualSlr,
            'error': None
        }

        date_fields = [
            'fecha_expedicion', 'vigencia_desde', 'vigencia_hasta', 'fecha_vencimiento',
            'fechaExpedicion', 'fechaInicialVigencia', 'fechaFinVigencia'
        ]

        failed_rows_data = []
        success_count = 0

        for index, row in df.iterrows():
            row_data_for_error = row.to_dict()
            row_data_for_error = {k: v for k, v in row_data_for_error.items() if pd.notna(v)}

            try:
                tipo_str = str(row.get('tipo_documento', '')).lower()
                if tipo_str == 'error':
                    tipo_str = str(row.get('document_type', '')).lower()
                
                placa = row.get('placa')

                if not all([tipo_str, placa]):
                    raise serializers.ValidationError("Faltan datos esenciales (tipo_documento, placa).")

                if tipo_str not in doc_serializers:
                     raise serializers.ValidationError(f"Tipo de documento '{tipo_str}' no es válido.")

                data = row.to_dict()
                for field in date_fields:
                    if field in data:
                        data[field] = normalize_date(data[field])

                if tipo_str == 'operacion':
                    data['fechaExpedicion'] = data.get('fecha_expedicion')
                    data['fechaInicialVigencia'] = data.get('vigencia_desde')
                    data['fechaFinVigencia'] = data.get('fecha_vencimiento')
                    data['numero'] = data.get('no_certificado')

                if tipo_str in ['contractual', 'extracontractual']:
                    data['fecha_inicio_vigencia'] = data.get('vigencia_desde')
                    data['fecha_fin_vigencia'] = data.get('vigencia_hasta')

                try:
                    vehiculo = Vehiculos.objects.get(placa=placa)
                except Vehiculos.DoesNotExist:
                    raise serializers.ValidationError(f"Vehículo con placa '{placa}' no encontrado.")
                
                soporte_path = None
                file_name = row.get('filename') 
                if file_name and request.FILES.get(file_name):
                    file_obj = request.FILES.get(file_name)
                    filename = f'docs_xlsx/{uuid.uuid4().hex}_{file_obj.name}'
                    soporte_path = default_storage.save(filename, file_obj)
                    data['soporte'] = soporte_path
                
                serializer_class = doc_serializers[tipo_str]
                serializer = serializer_class(data=data)

                serializer.is_valid(raise_exception=True)
                serializer.save(vehiculo=vehiculo)
                success_count += 1

            except Exception as e:
                error_detail = str(e.detail) if hasattr(e, 'detail') else str(e)
                row_data_for_error['motivo_error'] = error_detail
                failed_rows_data.append(row_data_for_error)
                continue
        
        if failed_rows_data:
            failed_df = pd.DataFrame(failed_rows_data)
            output_buffer = BytesIO()
            with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
                failed_df.to_excel(writer, index=False, sheet_name='Errores')
            output_buffer.seek(0)
            
            response = HttpResponse(
                output_buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename="reporte_errores_carga.xlsx"'
            response.status_code = status.HTTP_400_BAD_REQUEST 
            return response

        return Response(
            {'message': f'Carga masiva completada. Se crearon {success_count} registros exitosamente.'},
            status=status.HTTP_201_CREATED
        )